"""
ASI 360 — Universal Proposal Template Framework
================================================
SANDBOX ONLY — no BoldSign send, no PDF conversion.
Output: ASI360_Proposal_TEMPLATE_SANDBOX.docx

HOW TO USE THIS FILE:
1. Copy CLIENT_CONFIG and fill in all [FILL] values
2. Set proposal_type to "nvr_camera" or "access_control"
3. Set has_option_b to True/False
4. Run: python build_template.py
5. Review the DOCX output — then build the production version

FIELD INTELLIGENCE BRIEF — Questions to gather before filling in this template:
────────────────────────────────────────────────────────────────────────────────
IDENTITY
  □ Client first/last name and correct title (Owner / GM / IT Director / etc.)
  □ Company legal name AND trade name (if different)
  □ Full site address (street, city, state, zip)
  □ Client's direct email and phone
  □ Who else needs a copy? (CC list)

SITE WALK / DISCOVERY
  □ What is the site type? (bar, restaurant, office, retail, warehouse, residential)
  □ Describe the physical layout in 2–3 sentences
  □ What are the primary security concerns they verbalized?
  □ What is the existing system (brand/model) and what is wrong with it?
  □ How many cameras currently installed? Working vs. broken/offline?
  □ Camera generations (4K / 2K / 1080p / older) and brands present?
  □ Any specific camera locations they want addressed (entry, POS, back office, etc.)?
  □ NVR location: dedicated rack, shelf, closet — power available?
  □ Is existing cabling RG59 or Cat5/6? Any conduit runs needed?
  □ Is PoE switch adequate? (how many ports, PoE or PoE+ rated?)
  □ Do they have a managed network? Static IP scheme already in use?
  □ Remote access: who needs it? (owner, manager, off-site security?)
  □ What triggered this inquiry? (incident? system failure? audit? expansion?)

ACCESS CONTROL (if applicable)
  □ How many doors need to be controlled?
  □ Current lock type: key-only, deadbolt, electric strike, magloc?
  □ Credential preference: mobile app, fob, card, PIN pad?
  □ Who manages access? On-site or cloud-managed?
  □ Is there an existing access control system to integrate or replace?
  □ REX (request to exit) sensor needed on any doors?
  □ Door hours / schedule rules needed?

SCOPE CONFIRMATION
  □ Is this a swap (like-for-like) or an expansion?
  □ Any cameras to add beyond the existing count?
  □ Any cameras needing physical repair or replacement?
  □ Are there any out-of-scope items the client mentioned (acknowledge but exclude)?
  □ What is the client's timeline expectation?
  □ Have they gotten competing quotes? (affects discount strategy)

PRICING INPUTS
  □ Hardware costs from distributor quote (verify current pricing)
  □ Labor scope: is this a half-day, full-day, or multi-day install?
  □ Any special access fees, parking, lift equipment needed?
  □ What discount (if any) is appropriate for this client relationship?
  □ Is there a second option (e.g., more storage, additional camera)?

CLOSEOUT PLANNING
  □ Who signs? (is signer same as client contact, or separate decision-maker?)
  □ What is BoldSign signer email?
  □ Deposit method preference: check, ACH, card?
  □ Confirm install date window before writing into proposal
────────────────────────────────────────────────────────────────────────────────
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage
import io
import os

# ══════════════════════════════════════════════════════════════════════════════
# CLIENT CONFIG — Fill in ALL [FILL] values before running
# ══════════════════════════════════════════════════════════════════════════════

CLIENT_CONFIG = {

    # ── Identification ────────────────────────────────────────────────────────
    "proposal_number":  "QUO[FILL]",
    "proposal_type":    "nvr_camera",   # "nvr_camera"  |  "access_control"
    "date_str":         "[FILL — Month YYYY]",
    "valid_through":    "[FILL — Month DD, YYYY]",

    # ── Client ────────────────────────────────────────────────────────────────
    "client_name":      "[FILL — First Last]",
    "client_title":     "[FILL — Owner / General Manager / IT Director]",
    "client_company":   "[FILL — Company Name]",
    "client_address":   "[FILL — Street Address, City, CA ZIP]",
    "client_email":     "[FILL]@[FILL].com",

    # ── Site ──────────────────────────────────────────────────────────────────
    "site_name":        "[FILL — Venue / Property Name]",
    "site_address":     "[FILL — Street Address, City, CA ZIP]",
    "site_type":        "[FILL — bar / restaurant / law office / retail / warehouse]",
    "site_description": "[FILL — 1–2 sentence physical description of the site and its layout]",

    # ── Cover Letter ──────────────────────────────────────────────────────────
    "cover_pain":       "[FILL — Describe the specific operational problem they verbalized in discovery: "
                        "e.g., 'staff spending 4+ hours pulling footage', 'key fob system has no audit log', "
                        "'cameras keep going offline with no alert'. Use their words.]",

    "cover_approach":   "[FILL — 1–2 sentences describing ASI 360's approach for this project: "
                        "what is preserved, what is replaced, why it's minimally disruptive.]",

    "cover_scope_summary": "[FILL — 2–3 sentence scope summary: primary hardware, key labor items, "
                            "what client ends up with on completion day.]",

    # ── Needs Assessment ──────────────────────────────────────────────────────
    "needs": [
        # Format: ("Need Title", "Description of the specific pain in client's language")
        ("[FILL — Need #1 Title]",
         "[FILL — Describe this specific pain point. Be concrete: what happens today that shouldn't? "
         "What is the liability or operational impact?]"),

        ("[FILL — Need #2 Title]",
         "[FILL — Describe this specific pain point.]"),

        ("[FILL — Need #3 Title]",
         "[FILL — Describe this specific pain point.]"),

        # Add or remove needs as appropriate for this engagement
    ],

    "market_description": (
        "[FILL — Who is this client and what do they operate? What is at stake if security fails? "
        "e.g., 'Mad Oak Bar & Yard — a licensed food and beverage venue in Oakland's Jack London Square "
        "district. The venue operates indoor bar, dining, and outdoor yard areas...']"
    ),

    "solution_description": (
        "[FILL — What specifically is ASI 360 proposing to solve the identified needs? "
        "Name the primary hardware, key config work, and end-state. Be specific.]"
    ),

    "discovery_sources": [
        "[FILL — e.g., 'Site conversation and camera inventory review — Don Bucknor, ASI 360 COO, Month YYYY']",
        "[FILL — e.g., 'Site assessment photos — uploaded to Airtable intake, Month YYYY']",
        # Remove or add lines as needed
    ],

    # ── Scope of Work ─────────────────────────────────────────────────────────
    "scope_intro": (
        "[FILL — One sentence introducing the scope: "
        "e.g., 'Listed below are the specific measures ASI 360 will execute at [site_name] "
        "to upgrade the security video infrastructure.']"
    ),

    "scope_sections": [
        # Format: ("N. Section Title", ["bullet 1", "bullet 2", ...])
        ("1. [FILL — Primary Scope Item Title]", [
            "[FILL — specific action / deliverable]",
            "[FILL — specific action / deliverable]",
            "[FILL — specific action / deliverable]",
        ]),
        ("2. [FILL — Secondary Scope Item Title]", [
            "[FILL — specific action / deliverable]",
            "[FILL — specific action / deliverable]",
        ]),
        ("3. [FILL — Additional Scope Item]", [
            "[FILL — specific action / deliverable]",
        ]),
        # Typical NVR sections: Hardware Swap, Camera Integration, IP Scheme / Network,
        #   Remote Access, LumiSearch AI, Camera Refurbishment, Remote Access + Training
        # Typical Access Control sections: Hardware Install, Controller Config, Cloud Setup,
        #   Credential Enrollment, Door Schedule Programming, Training
    ],

    # ── Product Cards ─────────────────────────────────────────────────────────
    # Each entry: (image_path_or_None, img_width_inches, product_name, subtitle, [specs], value_text)
    # Set image_path to None for a placeholder gray box (no product image available yet)
    "product_cards": [
        (
            None,   # Replace with path like: "/path/to/product_image.jpg"
            1.8,
            "[FILL — Product Name / Model Number]",
            "[FILL — Product category: e.g., '32-Channel AI NVR' or 'Access Controller']",
            [
                "[FILL — Key spec: e.g., '32 channels, 1.5U rack, 4×SATA bays']",
                "[FILL — Key spec: e.g., 'LumiSearch+ AI search capability']",
                "[FILL — Key spec: e.g., 'Dual NIC, H.265+ compression']",
                "[FILL — Key spec: e.g., '3-year manufacturer warranty']",
            ],
            "[FILL — Value statement: why this product solves their specific need]",
        ),
        (
            None,
            1.8,
            "[FILL — Second Product Name / Model Number]",
            "[FILL — Product category]",
            [
                "[FILL — Key spec]",
                "[FILL — Key spec]",
                "[FILL — Key spec]",
            ],
            "[FILL — Value statement]",
        ),
        # Remove product_cards entry entirely if no products to feature in this proposal
    ],

    # ── Project Schedule ──────────────────────────────────────────────────────
    "schedule_intro": (
        "[FILL — One sentence framing the schedule: "
        "e.g., 'The following 5-phase schedule maps the project management framework to "
        "specific deliverables and dates for the [site_name] [project type].']"
    ),
    "lead_time_note": (
        "[FILL — Equipment lead time warning if applicable: "
        "e.g., 'NVR ships from distributor warehouse in Texas. Standard transit is 5 business days. "
        "Order placed [date] → estimated arrival [date]. No rush shipping required for [install date].']"
        "  —  Set to None to omit this note entirely."
    ),
    "install_date":         "[FILL — Day, Month DD, YYYY]",
    "install_backup_date":  "[FILL — Day, Month DD, YYYY]",

    "phases": [
        # Format: ("Phase N: Title  |  Date Range", [("N.N  Task Name", "task description"), ...])
        ("Phase 1: Conception & Initiation  |  [FILL — Date]", [
            ("1.1  Project Charter",
             "[FILL — What happens on contract execution: deposit, scope confirmation, project launch.]"),
            ("1.2  Plan Review",
             "[FILL — What site/network/infrastructure audit occurs: in-person, remote, Hub scan.]"),
            ("1.3  Initiation",
             "[FILL — What discovery/setup action kicks off the project: Hub deployment, IP audit, site photos.]"),
        ]),
        ("Phase 2: Definition & Planning  |  [FILL — Date Range]", [
            ("2.1  Scope & Goal Setting",
             "[FILL — Finalize camera count, storage option, IP scheme, or door count/credential type.]"),
            ("2.2  Budget Approval",
             "[FILL — Deposit receives hardware order. Lead time / shipping details.]"),
            ("2.3  Work Breakdown Schedule",
             "[FILL — Install date confirmed. Day-of task list finalized.]"),
            ("2.4  Project Schedule / Gantt",
             "[FILL — Equipment tracking. Client notified upon receipt.]"),
            ("2.5  Communication Plan",
             "[FILL — Update cadence. Who gets status updates and how (SMS, email).]"),
            ("2.6  Risk Management",
             "[FILL — What compatibility checks run pre-install. What is the contingency date.]"),
        ]),
        ("Phase 3: Launch & Execution  |  [FILL — Date Range]", [
            ("3.1  Status & Tracking",
             "[FILL — Equipment received, staged, firmware updated, pre-configured.]"),
            ("3.2  KPI Monitoring",
             "[FILL — Pre-install prep: IP table staged / credentials programmed / controller pre-configured.]"),
            ("3.3  Quality Assurance",
             "[FILL — Protocol compatibility verified / door hardware tested / reader tested.]"),
            ("3.4  Progress Forecasts",
             "[FILL — Final day-of schedule confirmed with client. Site access arranged.]"),
        ]),
        ("Phase 4: Performance & Control  |  [FILL — Install Date] — Installation Day", [
            ("4.1  Objective Execution",
             "[FILL — Primary install action: NVR swap / controller mount / door hardware install.]"),
            ("4.2  Quality Deliverables",
             "[FILL — What is verified live: all cameras on NVR / all doors responding / remote access confirmed.]"),
            ("4.3  Effort & Cost Tracking",
             "[FILL — Secondary deliverable completed: new camera installed / AI activated / mobile app enrolled.]"),
            ("4.4  Performance Review",
             "[FILL — Training completed: footage review / door access management / alert setup.]"),
        ]),
        ("Phase 5: Project Close  |  [FILL — Close Date Range]", [
            ("5.1  Postmortem",
             "[FILL — Punch list items addressed: broken camera diagnosed / door alignment adjusted / etc.]"),
            ("5.2  Project Punchlist",
             "[FILL — Documentation delivered: IP map / credential sheet / door schedule printout.]"),
            ("5.3  Final Report",
             "[FILL — Sign-off. Warranty activated. Balance invoice issued. Project closed in VTiger.]"),
        ]),
    ],

    # ── NVR / Camera-specific (only used when proposal_type == "nvr_camera") ──
    "nvr_config": {
        "current_nvr":       "[FILL — Brand and model of the system being replaced]",
        "camera_count":      "[FILL — XX–XX]",
        "ip_scheme":         "[FILL — 192.168.X.101–.XXX  (or confirm client's existing scheme)]",
        "has_fleet_roadmap": True,   # Set False to omit the Luminys fleet roadmap table
        "roadmap_rows": [
            # (phase_num, action, lumisearch_coverage, timeline)
            ("1", "[FILL — NVR swap + X new Luminys camera(s) (this proposal)]",   "[FILL — X camera(s) AI-native]", "Now"),
            ("2", "[FILL — Replace highest-priority cameras with Luminys units]",    "[FILL — X–X cameras AI-native]", "[FILL — 3–6 months]"),
            ("3", "[FILL — Replace remaining legacy cameras as they age out]",       "[FILL — X+ cameras AI-native]", "[FILL — 6–18 months]"),
            ("4", "[FILL — Full Luminys fleet — LumiSearch at maximum effectiveness]", "All cameras",               "[FILL — 18–24 months]"),
        ],
    },

    # ── Access Control-specific (only used when proposal_type == "access_control") ──
    "ac_config": {
        "current_system":   "[FILL — existing lock/badge system: key-only, legacy reader, none]",
        "door_count":       "[FILL — number]",
        "credential_type":  "[FILL — mobile / fob / card / PIN]",
        "cloud_platform":   "[FILL — Keri Borealis / Openpath / Brivo / standalone]",
        "notes":            "[FILL — any special conditions: REX sensors, door schedule rules, integration points]",
    },

    # ── Investment — Option A ─────────────────────────────────────────────────
    "has_option_b":   True,   # Set False to show only Option A
    "option_a_label": "Option A \u2014 [FILL — brief description, e.g., '2-Week Storage (8TB)']",
    "option_b_label": "Option B \u2014 [FILL — brief description]  \u2014  Recommended",
    "option_a_color": "003B71",   # Navy — default Option A color
    "option_b_color": "1A6E1A",   # Green — default Option B (Recommended) color

    "option_a": {
        "hw_items": [
            # (description, qty, unit_price, total)
            ("[FILL — Primary hardware: NVR / controller / reader / camera]",   "1", "$[FILL]", "$[FILL]"),
            ("[FILL — Storage / HDD / accessory]",                              "1", "$[FILL]", "$[FILL]"),
            ("[FILL — Additional hardware if applicable]",                      "1", "$[FILL]", "$[FILL]"),
            ("Cables, Connectors, Mounting Hardware (misc)",                    "1", "$[FILL]", "$[FILL]"),
        ],
        "labor_items": [
            # (description, qty, unit_price, total)
            ("Labor: [FILL — primary install labor description]",               "1", "$[FILL]", "$[FILL]"),
            ("Labor: [FILL — IP / network / config labor description]",         "1", "$[FILL]", "$[FILL]"),
            ("Labor: [FILL — remote access / training / closeout labor]",       "1", "$[FILL]", "$[FILL]"),
            # Add additional labor lines as needed
        ],
        "hw_sub":             "$[FILL]",
        "labor_sub":          "$[FILL]",
        "gross":              "$[FILL]",
        "discount":           "\u2212$[FILL]",
        "net":                "$[FILL]",
        "deposit_hw":         "$[FILL]",
        "deposit_labor":      "$[FILL]",
        "deposit_total":      "$[FILL]",
        "balance_labor":      "$[FILL]",
        "balance_discount":   "\u2212$[FILL]",
        "balance_total":      "$[FILL]",
    },

    # Set option_b to None if has_option_b is False
    "option_b": {
        "hw_items": [
            ("[FILL — Primary hardware (same or upgraded)]",                    "1", "$[FILL]", "$[FILL]"),
            ("[FILL — Upgraded storage / additional HDD]",                      "2", "$[FILL]", "$[FILL]"),
            ("[FILL — Additional hardware if Option B adds items]",             "1", "$[FILL]", "$[FILL]"),
            ("Cables, Connectors, Mounting Hardware (misc)",                    "1", "$[FILL]", "$[FILL]"),
        ],
        "labor_items": [
            ("Labor: [FILL — primary install labor description]",               "1", "$[FILL]", "$[FILL]"),
            ("Labor: [FILL — IP / network / config labor description]",         "1", "$[FILL]", "$[FILL]"),
            ("Labor: [FILL — Option B-specific labor, e.g., dual-drive RAID]",  "1", "$[FILL]", "$[FILL]"),
            ("Labor: [FILL — remote access / training / closeout labor]",       "1", "$[FILL]", "$[FILL]"),
        ],
        "hw_sub":             "$[FILL]",
        "labor_sub":          "$[FILL]",
        "gross":              "$[FILL]",
        "discount":           "\u2212$[FILL]",
        "net":                "$[FILL]",
        "deposit_hw":         "$[FILL]",
        "deposit_labor":      "$[FILL]",
        "deposit_total":      "$[FILL]",
        "balance_labor":      "$[FILL]",
        "balance_discount":   "\u2212$[FILL]",
        "balance_total":      "$[FILL]",
    },

    # ── Risk Analysis ─────────────────────────────────────────────────────────
    "risks": [
        # Format: ("Risk Title", "Analysis", "Resolution", "Contingency")
        (
            "Equipment Delivery Delay",
            "[FILL — Describe shipping source and standard lead time: "
            "e.g., 'Distributor ships from Texas — 5 business days standard.']",
            "Order confirmed same day contract signed. Carrier tracking shared with client upon shipment.",
            "[FILL — Backup install date and impact statement: e.g., 'Backup install date: [date]. No impact to project close.']",
        ),
        (
            "[FILL — Second Risk Title]",
            "[FILL — Describe the specific risk condition and what triggers it]",
            "[FILL — How ASI 360 prevents or mitigates this risk before it occurs]",
            "[FILL — What happens if it occurs anyway — fallback plan and cost impact]",
        ),
        (
            "Network Infrastructure Gaps",
            "[FILL — e.g., 'Client switch/router may not support PoE+ for new cameras or full channel throughput.']",
            "Audit switch specs during Phase 1 Hub discovery scan.",
            "[FILL — e.g., 'Additional PoE+ switch quoted separately if required (~$150–$300 depending on port count).']",
        ),
        (
            "Scope Expansion",
            "Client requests additional cameras, doors, or scope changes mid-project.",
            "Change order process documented in contract. All additions require written approval.",
            "[FILL — Change order rate: e.g., '$150/camera add-on labor rate. Hardware quoted at current pricing.']",
        ),
        (
            "[FILL — Project-specific risk, e.g., 'Legacy Camera Compatibility' or 'Door Hardware Clearance Issue']",
            "[FILL — Analysis]",
            "[FILL — Resolution]",
            "[FILL — Contingency]",
        ),
    ],

    # ── Deliverables ──────────────────────────────────────────────────────────
    "deliverables": [
        # Format: ("Deliverable Title", "Specific description of what is handed off")
        ("[FILL — Primary hardware deliverable]",
         "[FILL — e.g., 'Luminys R54-32NA live in rack — all XX cameras streaming on correct channels']"),
        ("Static IP Documentation",
         "[FILL — e.g., 'Complete IP map: camera name → static IP → NVR channel → location label']"),
        ("Credentials Document",
         "All usernames, passwords, and access levels for [FILL — NVR / controller / cloud portal / mobile app]"),
        ("Remote Access Configured",
         "[FILL — e.g., 'Web browser (LAN/WAN) + mobile app (iOS/Android) verified working off-site']"),
        ("[FILL — Additional hardware deliverable if applicable]",
         "[FILL — e.g., 'New camera installed, cabled, registered, AI activated']"),
        ("30-Minute Staff Training",
         "[FILL — What is covered in training: e.g., 'Footage review, clip export, mobile app, alert management — with client']"),
        ("90-Day Labor Warranty",
         "Covers all configured settings, IP assignments, remote access setup, and [FILL — any access control config]"),
        ("Manufacturer Hardware Warranty",
         "[FILL — Warranty terms for each hardware item: e.g., 'NVR: 3-year. Camera: 1-year. Controller: 2-year.']"),
        ("Work Order Closed",
         "Signed work order closed in VTiger. Project archived in ASI 360 Hub."),
    ],
}

# ── Paths ─────────────────────────────────────────────────────────────────────
LOGO    = "/Users/dbucknor/Downloads/artrends-attachments/ASI Logo/ASI_Logo-1_300P.jpg"
OUT     = "/Users/dbucknor/Downloads/360_Quotes_Engine/ASI360_Proposal_TEMPLATE_SANDBOX.docx"
PHONE   = "510-495-0905"

# ── Brand colors ──────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x00, 0x3B, 0x71)
TEAL  = RGBColor(0x00, 0x7A, 0x99)
GRAY  = RGBColor(0x55, 0x55, 0x55)
RED   = RGBColor(0xCC, 0x00, 0x00)
GREEN = RGBColor(0x1A, 0x6E, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS  (identical to build_v5.py)
# ══════════════════════════════════════════════════════════════════════════════

def img_buf(path, quality=92):
    img = PILImage.open(path).convert("RGB")
    buf = io.BytesIO()
    img.save(buf, "JPEG", quality=quality)
    buf.seek(0)
    return buf

def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="CCCCCC", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_cell_width(cell, inches):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(Inches(inches).pt * 20)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    run = p.add_run(text)
    if level == 1:
        run.font.color.rgb = NAVY
        run.font.size = Pt(14)
    elif level == 2:
        run.font.color.rgb = TEAL
        run.font.size = Pt(11)
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '80')
    sp.set(qn('w:after'), '40')
    pPr.append(sp)
    return p

def body(doc, text, italic=False, bold=False, color=None, size=10, align=None):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '40')
    pPr.append(sp)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p

def bullet(doc, text, size=10, color=None):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:hanging'), '180')
    pPr.append(ind)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '30')
    pPr.append(sp)
    run = p.add_run(f"\u2022  {text}")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def note(doc, text):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '40')
    sp.set(qn('w:after'), '40')
    pPr.append(sp)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RED
    run.bold = True
    return p

def spacer(doc, pts=4):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(pts * 20))
    sp.set(qn('w:after'), '0')
    pPr.append(sp)
    return p

def divider(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '003B71')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_image(doc, path, width_in, align=WD_ALIGN_PARAGRAPH.CENTER, caption=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.add_run().add_picture(img_buf(path), width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = GRAY
    return p

def placeholder_image_cell(cell, width_in, label="[PRODUCT IMAGE]"):
    """Gray placeholder box when no product image is available."""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = GRAY
    set_cell_shading(cell, "DDDDDD")

def product_card(doc, img_path, img_w, name, subtitle, specs, value_text):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = tbl.rows[0]
    left, right = row.cells[0], row.cells[1]
    set_cell_width(left, 2.3)
    set_cell_width(right, 4.1)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_border(left, "AAAAAA")
    set_cell_border(right, "AAAAAA")
    set_cell_shading(left, "F2F7FB")

    if img_path and os.path.exists(img_path):
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(img_buf(img_path), width=Inches(img_w))
    else:
        placeholder_image_cell(left, img_w)

    rp = right.paragraphs[0]
    nr = rp.add_run(name)
    nr.bold = True
    nr.font.size = Pt(11)
    nr.font.color.rgb = NAVY

    sub = right.add_paragraph()
    sr = sub.add_run(subtitle)
    sr.italic = True
    sr.font.size = Pt(9)
    sr.font.color.rgb = GRAY

    for spec in specs:
        sp = right.add_paragraph()
        pPr = sp._element.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '180')
        ind.set(qn('w:hanging'), '180')
        pPr.append(ind)
        sp.add_run(f"\u2022  {spec}").font.size = Pt(9)

    vp = right.add_paragraph()
    vr = vp.add_run(value_text)
    vr.italic = True
    vr.font.size = Pt(9)
    vr.font.color.rgb = GREEN

def section_header_bar(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, "003B71")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(12)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def page_break(doc):
    doc.add_page_break()

def investment_table(doc, opt_label, opt_color, hw_items, labor_items,
                     hw_sub, labor_sub, gross, discount, net,
                     deposit_hw, deposit_labor, deposit_total,
                     balance_labor, balance_discount, balance_total):
    # Option header bar
    hdr_tbl = doc.add_table(rows=1, cols=1)
    cell = hdr_tbl.rows[0].cells[0]
    set_cell_shading(cell, opt_color)
    r = cell.paragraphs[0].add_run(f"  {opt_label}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = WHITE
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    all_rows = []
    all_rows.append(("HEADER",  ["Line Item", "Qty", "Unit Price", "Total"]))
    all_rows.append(("SECTION", ["HARDWARE  \u2014  100% Due with Deposit Order", "", "", ""]))
    for item in hw_items:
        all_rows.append(("HW", list(item)))
    all_rows.append(("SUBTOTAL", ["", "", "HARDWARE SUBTOTAL", hw_sub]))
    all_rows.append(("SECTION",  ["LABOR & SERVICES  \u2014  50% Deposit / 50% on Completion", "", "", ""]))
    for item in labor_items:
        all_rows.append(("LABOR", list(item)))
    all_rows.append(("SUBTOTAL",   ["", "", "LABOR SUBTOTAL", labor_sub]))
    all_rows.append(("TOTAL_ROW",  ["", "", "GROSS TOTAL (pre-discount)", gross]))
    all_rows.append(("DISCOUNT",   ["", "", "ASI 360 Client Discount", discount]))
    all_rows.append(("NET",        ["", "", "NET TOTAL", net]))
    all_rows.append(("SECTION",    ["PAYMENT SCHEDULE", "", "", ""]))
    all_rows.append(("PAY_LABEL",  ["", "", "Deposit Invoice (due to schedule):", deposit_total]))
    all_rows.append(("PAY_DETAIL", ["    Hardware \u2014 100%", "", "", deposit_hw]))
    all_rows.append(("PAY_DETAIL", ["    Labor deposit \u2014 50%", "", "", deposit_labor]))
    all_rows.append(("PAY_LABEL",  ["", "", "Balance Due on Completion:", balance_total]))
    all_rows.append(("PAY_DETAIL", ["    Labor balance \u2014 50%", "", "", balance_labor]))
    all_rows.append(("PAY_DETAIL", ["    Less: ASI 360 client discount", "", "", balance_discount]))

    tbl = doc.add_table(rows=len(all_rows), cols=4)
    tbl.style = 'Table Grid'
    hw_idx = labor_idx = 0

    for ri, (row_type, values) in enumerate(all_rows):
        row = tbl.rows[ri]
        cells = row.cells

        if row_type == "HEADER":
            for ci, h in enumerate(values):
                set_cell_shading(cells[ci], "D6E4F0")
                r = cells[ci].paragraphs[0].add_run(h)
                r.bold = True
                r.font.size = Pt(9)

        elif row_type == "SECTION":
            for ci in range(4):
                set_cell_shading(cells[ci], "1A4B78")
            r = cells[0].paragraphs[0].add_run(values[0])
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = WHITE
            for ci in [1, 2, 3]:
                cells[ci].paragraphs[0].add_run("")

        elif row_type in ("HW", "LABOR"):
            bg = "F2F7FB" if (hw_idx if row_type == "HW" else labor_idx) % 2 == 0 else "FFFFFF"
            if row_type == "HW":
                hw_idx += 1
            else:
                labor_idx += 1
            for ci, val in enumerate(values):
                set_cell_shading(cells[ci], bg)
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(9)
                if ci in [2, 3]:
                    cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "SUBTOTAL":
            for ci, val in enumerate(values):
                set_cell_shading(cells[ci], "E8E8E8")
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(9)
                r.bold = True
                if ci in [2, 3]:
                    cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "TOTAL_ROW":
            for ci, val in enumerate(values):
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(9)
                if ci in [2, 3]:
                    cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "DISCOUNT":
            for ci, val in enumerate(values):
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(9)
                r.font.color.rgb = GREEN
                if ci in [2, 3]:
                    r.bold = True
                    cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "NET":
            for ci, val in enumerate(values):
                set_cell_shading(cells[ci], opt_color)
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(10)
                r.bold = True
                r.font.color.rgb = WHITE
                if ci in [2, 3]:
                    cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "PAY_LABEL":
            for ci, val in enumerate(values):
                set_cell_shading(cells[ci], "EEF4FB")
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(9)
                r.bold = True
                if ci in [2, 3]:
                    cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "PAY_DETAIL":
            for ci, val in enumerate(values):
                set_cell_shading(cells[ci], "F8F8F8")
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(8)
                r.italic = True
                if ci == 3:
                    cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ══════════════════════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════════════════════

C = CLIENT_CONFIG   # shorthand

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# Header
header = doc.sections[0].header
hdr_para = header.paragraphs[0]
hdr_para.clear()
hdr_run = hdr_para.add_run("ASI 360 — Allied Systems Integrations  |  Confidential Proposal")
hdr_run.font.size = Pt(8)
hdr_run.font.color.rgb = GRAY
pPr = hdr_para._element.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single')
bot.set(qn('w:sz'), '4')
bot.set(qn('w:color'), 'AAAAAA')
pBdr.append(bot)
pPr.append(pBdr)

# Footer
footer = doc.sections[0].footer
ftr_para = footer.paragraphs[0]
ftr_para.clear()
ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
ftr_run = ftr_para.add_run(f"ASI 360  |  {PHONE}  |  don@asi360.co  |  www.asi360.co")
ftr_run.font.size = Pt(8)
ftr_run.font.color.rgb = GRAY


# ── PAGE 1: COVER LETTER ──────────────────────────────────────────────────────
body(doc, C["date_str"])
spacer(doc, 6)
body(doc, C["client_name"])
body(doc, f"{C['client_title']}, {C['client_company']}")
body(doc, C["client_address"])
spacer(doc, 6)
body(doc, f"Dear {C['client_name'].split()[0]},")
spacer(doc, 4)
body(doc, C["cover_pain"])
spacer(doc, 4)
body(doc, C["cover_approach"])
spacer(doc, 4)
body(doc, C["cover_scope_summary"])
spacer(doc, 4)
body(doc, "This proposal is valid for 30 days from the date of issue. Questions? Reach me directly:")
spacer(doc, 6)
body(doc, "Sincerely,")
spacer(doc, 4)
body(doc, "Don Bucknor", bold=True)
body(doc, "COO / Founder — ASI 360 | Allied Systems Integrations")
body(doc, f"{PHONE}  |  don@asi360.co  |  www.asi360.co")


# ── PAGE 2: TITLE PAGE ────────────────────────────────────────────────────────
page_break(doc)

p_co = doc.add_paragraph()
p_co.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p_co.add_run(f"ASI 360 | Allied Systems Integrations\n{PHONE} | don@asi360.co | www.asi360.co")
r.font.size = Pt(9)
r.font.color.rgb = GRAY

spacer(doc, 4)
add_image(doc, LOGO, 3.0)
spacer(doc, 2)
divider(doc)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("Security System Proposal")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run(f"{C['client_company']}  |  {C['site_address']}")
r.font.size = Pt(12)
r.font.color.rgb = TEAL

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_date.add_run(f"{C['date_str']}  |  Prepared by: Don Bucknor, COO")
r.font.size = Pt(10)
r.font.color.rgb = GRAY
r.italic = True

divider(doc)
spacer(doc, 4)

# Description box
tbl = doc.add_table(rows=1, cols=1)
cell = tbl.rows[0].cells[0]
set_cell_shading(cell, "F2F7FB")
set_cell_border(cell, "003B71", "6")
p = cell.paragraphs[0]
r = p.add_run(C["cover_scope_summary"])
r.font.size = Pt(10)
p2 = cell.add_paragraph()
r2 = p2.add_run(f"Proposal Number: {C['proposal_number']}  |  Valid through: {C['valid_through']}")
r2.font.size = Pt(9)
r2.italic = True
r2.font.color.rgb = GRAY
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── PAGE 3: TABLE OF CONTENTS ─────────────────────────────────────────────────
page_break(doc)
section_header_bar(doc, "TABLE OF CONTENTS")
spacer(doc, 6)

toc_items = [
    ("Cover Letter", "1"),
    ("Title Page", "2"),
    ("Table of Contents", "3"),
    ("Needs Assessment", "4"),
    ("Scope of Work", "5"),
    ("Project Schedule & Milestones", "9"),
    ("Investment Summary", "11"),
    ("Risk Analysis", "13"),
    ("Project Deliverables", "14"),
    ("About ASI 360", "15"),
    ("Terms & Conditions", "16"),
    ("Authorization", "17"),
]

tbl_toc = doc.add_table(rows=len(toc_items), cols=2)
tbl_toc.style = 'Table Grid'
for ri, (name, pg) in enumerate(toc_items):
    row = tbl_toc.rows[ri]
    lc, rc = row.cells[0], row.cells[1]
    bg = "F2F7FB" if ri % 2 == 0 else "FFFFFF"
    set_cell_shading(lc, bg)
    set_cell_shading(rc, bg)
    set_cell_border(lc, "CCCCCC", "2")
    set_cell_border(rc, "CCCCCC", "2")
    lc.paragraphs[0].add_run(name).font.size = Pt(10)
    rr = rc.paragraphs[0].add_run(pg)
    rr.font.size = Pt(10)
    rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Quick Facts callout
spacer(doc, 10)
tbl_qf = doc.add_table(rows=1, cols=1)
cell_qf = tbl_qf.rows[0].cells[0]
set_cell_shading(cell_qf, "EEF4FB")
set_cell_border(cell_qf, "003B71", "6")
r_qft = cell_qf.paragraphs[0].add_run("  PROPOSAL QUICK FACTS")
r_qft.bold = True
r_qft.font.size = Pt(10)
r_qft.font.color.rgb = NAVY

opt_a = C["option_a"]
opt_b = C.get("option_b")

quick_facts = [
    ("Client",       f"{C['client_name']} \u2014 {C['client_company']}, {C['site_address']}"),
    ("Proposal #",   f"{C['proposal_number']}  |  Valid through {C['valid_through']}"),
    ("Option A",     f"{opt_a['net']} net  \u2022  Deposit: {opt_a['deposit_total']}  \u2022  Balance: {opt_a['balance_total']}  \u2022  {C['option_a_label'].split('\u2014')[1].strip() if '\u2014' in C['option_a_label'] else ''}"),
]
if C["has_option_b"] and opt_b:
    quick_facts.append(("Option B",
        f"{opt_b['net']} net  \u2022  Deposit: {opt_b['deposit_total']}  \u2022  Balance: {opt_b['balance_total']}  \u2022  {C['option_b_label'].split('\u2014')[1].strip() if '\u2014' in C['option_b_label'] else ''}  \u2022  Recommended"))
quick_facts += [
    ("Hardware",     "100% due with deposit  \u2022  Confirm lead time with distributor at time of order"),
    ("Install Date", C["install_date"]),
    ("Warranty",     "90-day labor  \u2022  Manufacturer hardware warranty per product"),
]
for label, val in quick_facts:
    fp = cell_qf.add_paragraph()
    r1 = fp.add_run(f"  {label}:  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = TEAL
    fp.add_run(val).font.size = Pt(9)
cell_qf.add_paragraph()


# ── PAGE 4: NEEDS ASSESSMENT ──────────────────────────────────────────────────
page_break(doc)
section_header_bar(doc, "NEEDS ASSESSMENT")
spacer(doc, 4)

body(doc, f"ASI 360 has identified the following security infrastructure needs at {C['site_name']} "
         f"which are not currently being met by the existing {C['nvr_config']['current_nvr'] if C['proposal_type'] == 'nvr_camera' else C['ac_config']['current_system']}.")
spacer(doc, 4)
body(doc, "Needs:", bold=True)

for i, (title, desc) in enumerate(C["needs"], 1):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '40')
    pPr.append(sp)
    r1 = p.add_run(f"{i}.  {title} \u2014 ")
    r1.bold = True
    r1.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)

spacer(doc, 4)
body(doc, "Market:", bold=True)
body(doc, C["market_description"])

spacer(doc, 4)
body(doc, "Solution:", bold=True)
body(doc, C["solution_description"])

spacer(doc, 4)
body(doc, "Sources:", bold=True)
for src in C["discovery_sources"]:
    body(doc, f"\u2022 {src}")


# ── PAGE 5+: SCOPE OF WORK ────────────────────────────────────────────────────
page_break(doc)
section_header_bar(doc, "SCOPE OF WORK")
spacer(doc, 4)
body(doc, C["scope_intro"])
spacer(doc, 2)

for section_title, bullets in C["scope_sections"]:
    heading(doc, section_title, 2)
    for b in bullets:
        bullet(doc, b)
    spacer(doc, 2)

# Product Cards (if any defined)
if C.get("product_cards"):
    spacer(doc, 4)
    heading(doc, "Equipment Specifications", 2)
    for card in C["product_cards"]:
        img_path, img_w, name, subtitle, specs, value_text = card
        product_card(doc, img_path, img_w, name, subtitle, specs, value_text)
        spacer(doc, 4)


# ── PROJECT SCHEDULE ──────────────────────────────────────────────────────────
page_break(doc)
section_header_bar(doc, "PROJECT SCHEDULE & MILESTONES")
spacer(doc, 4)
body(doc, C["schedule_intro"])
spacer(doc, 2)

if C.get("lead_time_note") and "[FILL" not in C["lead_time_note"] and C["lead_time_note"] != "None":
    note(doc, f"\u26a0  {C['lead_time_note']}")
spacer(doc, 4)

for phase_title, tasks in C["phases"]:
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '60')
    sp.set(qn('w:after'), '20')
    pPr.append(sp)
    r = p.add_run(phase_title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    for task_name, task_desc in tasks:
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '360')
        pPr.append(ind)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '30')
        pPr.append(sp)
        r1 = p.add_run(f"{task_name}  ")
        r1.bold = True
        r1.font.size = Pt(10)
        p.add_run(task_desc).font.size = Pt(10)

# Fleet Roadmap (NVR proposals only)
if C["proposal_type"] == "nvr_camera" and C["nvr_config"].get("has_fleet_roadmap"):
    spacer(doc, 8)
    heading(doc, "Luminys Fleet Transition Roadmap", 2)
    body(doc, "The goal is full Luminys unification for complete LumiSearch AI capability. "
              "This happens in natural replacement cycles \u2014 no forced early swaps.")
    spacer(doc, 2)

    roadmap_rows = C["nvr_config"]["roadmap_rows"]
    tbl_rm = doc.add_table(rows=len(roadmap_rows) + 1, cols=4)
    tbl_rm.style = 'Table Grid'
    for ci, h in enumerate(["Phase", "Action", "LumiSearch Coverage", "Timeline"]):
        cell = tbl_rm.rows[0].cells[ci]
        set_cell_shading(cell, "003B71")
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
    for ri, row_data in enumerate(roadmap_rows):
        for ci, val in enumerate(row_data):
            cell = tbl_rm.rows[ri + 1].cells[ci]
            cell.paragraphs[0].add_run(val).font.size = Pt(9)
            if ri % 2 == 0:
                set_cell_shading(cell, "F2F7FB")


# ── INVESTMENT SUMMARY ────────────────────────────────────────────────────────
page_break(doc)
section_header_bar(doc, "INVESTMENT SUMMARY")
spacer(doc, 4)

oa = C["option_a"]
investment_table(doc,
    opt_label=C["option_a_label"],
    opt_color=C["option_a_color"],
    hw_items=oa["hw_items"],
    labor_items=oa["labor_items"],
    hw_sub=oa["hw_sub"],
    labor_sub=oa["labor_sub"],
    gross=oa["gross"],
    discount=oa["discount"],
    net=oa["net"],
    deposit_hw=oa["deposit_hw"],
    deposit_labor=oa["deposit_labor"],
    deposit_total=oa["deposit_total"],
    balance_labor=oa["balance_labor"],
    balance_discount=oa["balance_discount"],
    balance_total=oa["balance_total"],
)

if C["has_option_b"] and C.get("option_b"):
    spacer(doc, 8)
    ob = C["option_b"]
    investment_table(doc,
        opt_label=C["option_b_label"],
        opt_color=C["option_b_color"],
        hw_items=ob["hw_items"],
        labor_items=ob["labor_items"],
        hw_sub=ob["hw_sub"],
        labor_sub=ob["labor_sub"],
        gross=ob["gross"],
        discount=ob["discount"],
        net=ob["net"],
        deposit_hw=ob["deposit_hw"],
        deposit_labor=ob["deposit_labor"],
        deposit_total=ob["deposit_total"],
        balance_labor=ob["balance_labor"],
        balance_discount=ob["balance_discount"],
        balance_total=ob["balance_total"],
    )

spacer(doc, 4)
body(doc, "Hardware is invoiced 100% with the deposit order. Labor is split 50% with the deposit "
         "and 50% on project completion. The ASI 360 client discount is applied against the balance "
         "due at completion.", size=9, italic=True)


# ── RISK ANALYSIS ─────────────────────────────────────────────────────────────
page_break(doc)
section_header_bar(doc, "RISK ANALYSIS")
spacer(doc, 4)
body(doc, "Our analysis of the requirements for this project identifies the following potential risks. "
         "This list is not exhaustive, but covers all primary risks identified during scope review.")
spacer(doc, 4)

risks = C["risks"]
tbl = doc.add_table(rows=len(risks) + 1, cols=4)
tbl.style = 'Table Grid'
for ci, h in enumerate(["Risk", "Analysis", "Resolution", "Contingency"]):
    cell = tbl.rows[0].cells[ci]
    set_cell_shading(cell, "003B71")
    r = cell.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = WHITE
for ri, (risk, analysis, resolution, contingency) in enumerate(risks):
    row = tbl.rows[ri + 1]
    for ci, val in enumerate([risk, analysis, resolution, contingency]):
        r = row.cells[ci].paragraphs[0].add_run(val)
        r.font.size = Pt(8)
        if ri % 2 == 0:
            set_cell_shading(row.cells[ci], "F2F7FB")


# ── DELIVERABLES ──────────────────────────────────────────────────────────────
spacer(doc, 8)
section_header_bar(doc, "PROJECT DELIVERABLES")
spacer(doc, 4)
body(doc, "The following is a complete list of all project deliverables upon project completion and client sign-off.")
spacer(doc, 2)

for i, (title, desc) in enumerate(C["deliverables"], 1):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '30')
    pPr.append(sp)
    r1 = p.add_run(f"{i}.  {title} \u2014 ")
    r1.bold = True
    r1.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)


# ── ABOUT ASI 360 ─────────────────────────────────────────────────────────────
page_break(doc)
section_header_bar(doc, "ABOUT ASI 360 \u2014 ALLIED SYSTEMS INTEGRATIONS")
spacer(doc, 4)
add_image(doc, LOGO, 2.0, align=WD_ALIGN_PARAGRAPH.LEFT)
spacer(doc, 2)

body(doc, "About Us", bold=True)
body(doc, "ASI 360 (Allied Systems Integrations) is an Oakland-based security systems integrator "
         "specializing in physical security infrastructure for small businesses, bars, restaurants, "
         "retail, and commercial venues throughout the Bay Area. We design, install, and support "
         "complete security ecosystems \u2014 from single-camera upgrades to full multi-site deployments "
         "\u2014 with a focus on practical, cost-effective solutions built on enterprise-grade hardware.")
spacer(doc, 2)
body(doc, "Our approach is always infrastructure-first: we audit what you have, preserve what works, "
         "and replace only what\u2019s holding you back. Every project closes with a complete IP map, "
         "credential document, and staff training session \u2014 no black-box systems that only we can service.")
spacer(doc, 4)

body(doc, "Services", bold=True)
for svc in [
    "Security NVR/DVR upgrades and hybrid camera integration",
    "IP camera installation (indoor/outdoor, fixed/PTZ, AI-capable)",
    "Static IP scheme design and full network documentation",
    "Remote access configuration (web + mobile) and alert setup",
    "LumiSearch AI deployment and phased fleet transitions",
    "Access control systems (Keri Systems \u2014 door readers, credentials, audit logs)",
    "ASI 360 Hub network monitoring \u2014 camera health, uptime alerts, remote diagnostics",
    "Service agreements and 90-day labor warranty on all installs",
]:
    bullet(doc, svc)

spacer(doc, 4)
body(doc, "Technology Partners", bold=True)
body(doc, "ASI 360 is an authorized reseller and integration partner for Luminys Systems (AI cameras "
         "and NVRs) and Keri Systems (access control). All hardware is sourced through authorized "
         "distributors \u2014 ensuring authentic, warranted equipment at competitive pricing.")

spacer(doc, 4)
body(doc, "How to Contact ASI 360", bold=True)
body(doc, "ASI 360 | Allied Systems Integrations")
body(doc, "Don Bucknor, COO/Founder")
body(doc, f"{PHONE}  |  don@asi360.co  |  www.asi360.co")
body(doc, "Oakland, CA")


# ── TERMS & CONDITIONS ────────────────────────────────────────────────────────
spacer(doc, 8)
section_header_bar(doc, "TERMS & CONDITIONS")
spacer(doc, 4)

terms = [
    "Hardware costs are due 100% with the deposit to place the equipment order. "
    "Labor is billed 50% with the deposit and 50% upon project completion and client sign-off.",
    "The ASI 360 client discount is applied against the labor balance due at project completion.",
    "90-day labor warranty on all configured settings, IP assignments, and remote access setup",
    "Hardware manufacturer warranty applies per product (varies by model \u2014 see Deliverables section)",
    "Camera refurbishment / door hardware repair is scoped separately \u2014 quoted after on-site diagnosis",
    "Change orders require written approval before execution",
    "Client responsible for providing site access and electrical power",
    "ASI 360 not responsible for pre-existing infrastructure failures discovered during install",
    "This proposal is valid for 30 days from date of issue",
    "All pricing subject to distributor availability at time of order",
]
for t in terms:
    bullet(doc, t)

spacer(doc, 6)
body(doc, "ASI 360 Hub Monitoring (Optional Post-Project Service)", bold=True)
body(doc, "The ASI 360 Hub is a lightweight network agent deployed during Phase 1 discovery. "
         "After project close, optional ongoing monitoring is available:", size=10)
for hub_term in [
    "Camera health monitoring \u2014 uptime alerts and offline camera detection pushed to your phone",
    "Remote diagnostics \u2014 ASI 360 technician access via secure VPN for troubleshooting without an on-site visit",
    "Power cycle management and Grafana uptime dashboard view",
    "First 6 months included at no charge. Ongoing: $25/month (month-to-month, cancel anytime)",
    "Hub monitoring is separate from the 90-day labor warranty and does not replace manufacturer warranties",
]:
    bullet(doc, hub_term, size=9)


# ── AUTHORIZATION ─────────────────────────────────────────────────────────────
page_break(doc)
section_header_bar(doc, "AUTHORIZATION")
spacer(doc, 4)
body(doc, "By signing below, Client authorizes ASI 360 / Allied Systems Integrations to proceed "
         "with the work described in this proposal under the terms and conditions stated herein.")
spacer(doc, 4)

# Option selection checkboxes
p_opt = doc.add_paragraph()
r1 = p_opt.add_run("Selected Option:  ")
r1.bold = True
r1.font.size = Pt(10)
p_opt.add_run(
    f"\u2610 {C['option_a_label'].split('\u2014')[0].strip()}  "
    f"\u2014 Net Total: {C['option_a']['net']}  "
    f"(Deposit: {C['option_a']['deposit_total']} / Balance: {C['option_a']['balance_total']})"
).font.size = Pt(10)

if C["has_option_b"] and C.get("option_b"):
    p_opt2 = doc.add_paragraph()
    p_opt2.add_run("                         ")
    p_opt2.add_run(
        f"\u2610 {C['option_b_label'].split('\u2014')[0].strip()}  "
        f"\u2014 Net Total: {C['option_b']['net']}  "
        f"(Deposit: {C['option_b']['deposit_total']} / Balance: {C['option_b']['balance_total']})"
    ).font.size = Pt(10)

spacer(doc, 6)

# Signature table
sig_tbl = doc.add_table(rows=3, cols=2)
sig_data = [
    ("Client Signature", "ASI 360 / Don Bucknor, COO"),
    (f"Printed Name — {C['client_name']}", "Don Bucknor"),
    ("Date", "Date"),
]
for ri, (left_label, right_label) in enumerate(sig_data):
    lc = sig_tbl.rows[ri].cells[0]
    rc = sig_tbl.rows[ri].cells[1]
    set_cell_border(lc, "003B71")
    set_cell_border(rc, "003B71")
    lr = lc.paragraphs[0].add_run(left_label)
    lr.font.size = Pt(9)
    lr.font.color.rgb = TEAL
    lr.bold = True
    rr = rc.paragraphs[0].add_run(right_label)
    rr.font.size = Pt(9)
    rr.font.color.rgb = TEAL
    rr.bold = True
    lc.add_paragraph()
    rc.add_paragraph()

spacer(doc, 6)
# BoldSign tags (white text — invisible in print, used by BoldSign text-tag mode)
tag_p = doc.add_paragraph()
tag_run = tag_p.add_run(
    "{{sign|1|*|Client Signature|sig_client}} "
    "{{date|1|*|Date Signed|date_client}} "
    "{{text|1|*|Printed Name|name_client}}"
)
tag_run.font.color.rgb = WHITE
tag_run.font.size = Pt(8)

tag_p2 = doc.add_paragraph()
tag_run2 = tag_p2.add_run(
    "{{sign|2|*|ASI 360 Authorized Signature|sig_asi360}} "
    "{{date|2|*|Date|date_asi360}}"
)
tag_run2.font.color.rgb = WHITE
tag_run2.font.size = Pt(8)

spacer(doc, 4)
body(doc, f"Questions? Contact Don Bucknor  |  {PHONE}  |  don@asi360.co  |  www.asi360.co",
     size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)


# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"Saved: {OUT}")
