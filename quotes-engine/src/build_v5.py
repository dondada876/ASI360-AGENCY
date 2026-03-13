from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage
import io

# ── Paths ──────────────────────────────────────────────────────────────────────
IMG_DIR  = "/Users/dbucknor/Downloads/360_Quotes_Engine/luminys_images"
LOGO     = "/Users/dbucknor/Downloads/artrends-attachments/ASI Logo/ASI_Logo-1_300P.jpg"
OUT_DOCX = "/Users/dbucknor/Downloads/360_Quotes_Engine/MadOak_NVR_Upgrade_Proposal_v5_March2026.docx"

PHONE = "510-495-0905"

# ── Brand colors ───────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x00, 0x3B, 0x71)
TEAL  = RGBColor(0x00, 0x7A, 0x99)
GRAY  = RGBColor(0x55, 0x55, 0x55)
RED   = RGBColor(0xCC, 0x00, 0x00)
GREEN = RGBColor(0x1A, 0x6E, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = "D6E4F0"

# ── Helpers ────────────────────────────────────────────────────────────────────
def img_buf(path, quality=92):
    img = PILImage.open(path).convert("RGB")
    buf = io.BytesIO()
    img.save(buf, "JPEG", quality=quality)
    buf.seek(0)
    return buf

def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="CCCCCC", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_cell_width(cell, inches):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(Inches(inches).pt * 20)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    run = p.add_run(text)
    if level == 1:
        run.font.color.rgb = NAVY
        run.font.size = Pt(14)
    elif level == 2:
        run.font.color.rgb = TEAL
        run.font.size = Pt(11)
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '80')
    sp.set(qn('w:after'), '40')
    pPr.append(sp)
    return p

def body(doc, text, italic=False, bold=False, color=None, size=10, align=None):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '40')
    pPr.append(sp)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    return p

def bullet(doc, text, size=10, color=None):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:hanging'), '180')
    pPr.append(ind)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '30')
    pPr.append(sp)
    run = p.add_run(f"\u2022  {text}")
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return p

def note(doc, text):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '40')
    sp.set(qn('w:after'), '40')
    pPr.append(sp)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RED
    run.bold = True
    return p

def spacer(doc, pts=4):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(pts * 20))
    sp.set(qn('w:after'), '0')
    pPr.append(sp)
    return p

def divider(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '003B71')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_image(doc, path, width_in, align=WD_ALIGN_PARAGRAPH.CENTER, caption=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.add_run().add_picture(img_buf(path), width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = GRAY
    return p

def product_card(doc, img_path, img_w, name, subtitle, specs, value_text):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = tbl.rows[0]
    left, right = row.cells[0], row.cells[1]
    set_cell_width(left, 2.3); set_cell_width(right, 4.1)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_border(left, "AAAAAA"); set_cell_border(right, "AAAAAA")
    set_cell_shading(left, "F2F7FB")

    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run().add_picture(img_buf(img_path), width=Inches(img_w))

    rp = right.paragraphs[0]
    nr = rp.add_run(name)
    nr.bold = True; nr.font.size = Pt(11); nr.font.color.rgb = NAVY

    sub = right.add_paragraph()
    sr = sub.add_run(subtitle)
    sr.italic = True; sr.font.size = Pt(9); sr.font.color.rgb = GRAY

    for spec in specs:
        sp = right.add_paragraph()
        pPr = sp._element.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '180'); ind.set(qn('w:hanging'), '180')
        pPr.append(ind)
        sp.add_run(f"\u2022  {spec}").font.size = Pt(9)

    vp = right.add_paragraph()
    vr = vp.add_run(value_text)
    vr.italic = True; vr.font.size = Pt(9); vr.font.color.rgb = GREEN

def section_header_bar(doc, text):
    """Full-width navy bar with white section title."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, "003B71")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(12)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def section_sub_header(doc, text, bg_color="1A4B78"):
    """Narrower sub-section bar (hardware / labor divider)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, bg_color)
    p = cell.paragraphs[0]
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(9)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def page_break(doc):
    doc.add_page_break()

def investment_table(doc, opt_label, opt_color, hw_items, labor_items,
                     hw_sub, labor_sub, gross, discount, net,
                     deposit_hw, deposit_labor, deposit_total,
                     balance_labor, balance_discount, balance_total):
    """Build a full investment table with hardware/labor split and payment breakdown."""
    # Option header
    hdr_tbl = doc.add_table(rows=1, cols=1)
    cell = hdr_tbl.rows[0].cells[0]
    set_cell_shading(cell, opt_color)
    r = cell.paragraphs[0].add_run(f"  {opt_label}")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = WHITE
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # Main line-item table
    all_rows = []

    # Column header
    all_rows.append(("HEADER", ["Line Item", "Qty", "Unit Price", "Total"]))

    # Hardware sub-header
    all_rows.append(("SECTION", ["HARDWARE  \u2014  100% Due with Deposit Order", "", "", ""]))

    for item in hw_items:
        all_rows.append(("HW", list(item)))

    all_rows.append(("SUBTOTAL", ["", "", "HARDWARE SUBTOTAL", hw_sub]))

    # Labor sub-header
    all_rows.append(("SECTION", ["LABOR & SERVICES  \u2014  50% Deposit / 50% on Completion", "", "", ""]))

    for item in labor_items:
        all_rows.append(("LABOR", list(item)))

    all_rows.append(("SUBTOTAL", ["", "", "LABOR SUBTOTAL", labor_sub]))

    # Totals
    all_rows.append(("TOTAL_ROW", ["", "", "GROSS TOTAL (pre-discount)", gross]))
    all_rows.append(("DISCOUNT", ["", "", "ASI 360 Client Discount", discount]))
    all_rows.append(("NET", ["", "", "NET TOTAL", net]))

    # Payment breakdown sub-header
    all_rows.append(("SECTION", ["PAYMENT SCHEDULE", "", "", ""]))
    all_rows.append(("PAY_LABEL", ["", "", "Deposit Invoice (due to schedule):", deposit_total]))
    all_rows.append(("PAY_DETAIL", [f"    Hardware \u2014 100%", "", "", deposit_hw]))
    all_rows.append(("PAY_DETAIL", [f"    Labor deposit \u2014 50%", "", "", deposit_labor]))
    all_rows.append(("PAY_LABEL", ["", "", "Balance Due on Completion:", balance_total]))
    all_rows.append(("PAY_DETAIL", ["    Labor balance \u2014 50%", "", "", balance_labor]))
    all_rows.append(("PAY_DETAIL", ["    Less: ASI 360 client discount", "", "", balance_discount]))

    tbl = doc.add_table(rows=len(all_rows), cols=4)
    tbl.style = 'Table Grid'

    hw_idx = 0
    labor_idx = 0

    for ri, (row_type, values) in enumerate(all_rows):
        row = tbl.rows[ri]
        cells = row.cells

        if row_type == "HEADER":
            for ci, h in enumerate(values):
                set_cell_shading(cells[ci], "D6E4F0")
                r = cells[ci].paragraphs[0].add_run(h)
                r.bold = True; r.font.size = Pt(9)

        elif row_type == "SECTION":
            # Merge all 4 cells visually with shading
            set_cell_shading(cells[0], "1A4B78")
            set_cell_shading(cells[1], "1A4B78")
            set_cell_shading(cells[2], "1A4B78")
            set_cell_shading(cells[3], "1A4B78")
            r = cells[0].paragraphs[0].add_run(values[0])
            r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
            for ci in [1,2,3]:
                cells[ci].paragraphs[0].add_run("")

        elif row_type in ("HW", "LABOR"):
            bg = "F2F7FB" if (hw_idx if row_type == "HW" else labor_idx) % 2 == 0 else "FFFFFF"
            if row_type == "HW": hw_idx += 1
            else: labor_idx += 1
            for ci, val in enumerate(values):
                set_cell_shading(cells[ci], bg)
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(9)
                if ci in [2, 3]: cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "SUBTOTAL":
            for ci, val in enumerate(values):
                set_cell_shading(cells[ci], "E8E8E8")
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(9); r.bold = True
                if ci in [2, 3]: cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "TOTAL_ROW":
            for ci, val in enumerate(values):
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(9)
                if ci in [2, 3]: cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "DISCOUNT":
            for ci, val in enumerate(values):
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(9); r.font.color.rgb = GREEN
                if ci in [2, 3]:
                    r.bold = True
                    cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "NET":
            for ci, val in enumerate(values):
                set_cell_shading(cells[ci], opt_color)
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(10); r.bold = True; r.font.color.rgb = WHITE
                if ci in [2, 3]: cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "PAY_LABEL":
            for ci, val in enumerate(values):
                set_cell_shading(cells[ci], "EEF4FB")
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(9); r.bold = True
                if ci in [2, 3]: cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif row_type == "PAY_DETAIL":
            for ci, val in enumerate(values):
                set_cell_shading(cells[ci], "F8F8F8")
                r = cells[ci].paragraphs[0].add_run(val)
                r.font.size = Pt(8); r.italic = True
                if ci == 3: cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ══════════════════════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Set page margins
for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# ── Header ────────────────────────────────────────────────────────────────────
header = doc.sections[0].header
hdr_para = header.paragraphs[0]
hdr_para.clear()
hdr_run = hdr_para.add_run("ASI 360 — Allied Systems Integrations  |  Confidential Proposal")
hdr_run.font.size = Pt(8)
hdr_run.font.color.rgb = GRAY
pPr = hdr_para._element.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4'); bot.set(qn('w:color'), 'AAAAAA')
pBdr.append(bot); pPr.append(pBdr)

footer = doc.sections[0].footer
ftr_para = footer.paragraphs[0]
ftr_para.clear()
ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
ftr_run = ftr_para.add_run(f"ASI 360  |  {PHONE}  |  don@asi360.co  |  asi360.co")
ftr_run.font.size = Pt(8)
ftr_run.font.color.rgb = GRAY

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER LETTER
# ══════════════════════════════════════════════════════════════════════════════
body(doc, "March 6, 2026")
spacer(doc, 6)
body(doc, "Daniel Cukierman")
body(doc, "Owner, Mad Oak Bar & Yard")
body(doc, "135 12th Street, Oakland, CA 94607")
spacer(doc, 6)
body(doc, "Dear Daniel,")
spacer(doc, 4)
body(doc, "Thank you for the conversation about your current security setup at Mad Oak. The situation you described — staff spending 4+ hours locating and exporting footage after incidents — is exactly the kind of operational liability that a modern NVR upgrade resolves. When you can't pull footage quickly, you lose the window for police assistance, insurance documentation, and internal accountability. That's a real business risk for a venue running a licensed bar.")
spacer(doc, 4)
body(doc, "ASI 360 specializes in security infrastructure transitions for small businesses and venues. Our approach is straightforward: we preserve everything that's already working — your existing cameras, all your cabling, and your current rack infrastructure — while replacing the aging system that's holding you back. No wall cuts, no new cable runs, no disruption to your operations beyond a single install day.")
spacer(doc, 4)
body(doc, "This proposal (QUO202355) covers a full NVR swap to the Luminys R54-32NA 32-channel AI platform, one new 4K LumiSearch AI camera for the dining area entry, unified static IP management for all 22–24 cameras, full remote access setup, and a 30-minute staff training session. We've also included a phased roadmap showing how the AI search capability expands as your camera fleet naturally upgrades over the next 18–24 months.")
spacer(doc, 4)
body(doc, "This proposal is valid for 30 days from the date of issue. Questions? Reach me directly:")
spacer(doc, 6)
body(doc, "Sincerely,")
spacer(doc, 4)
body(doc, "Don Bucknor", bold=True)
body(doc, "COO / Founder — ASI 360 | Allied Systems Integrations")
body(doc, f"{PHONE}  |  don@asi360.co  |  asi360.co")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

# Company info top right
p_co = doc.add_paragraph()
p_co.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p_co.add_run(f"ASI 360 | Allied Systems Integrations\n{PHONE} | don@asi360.co | asi360.co")
r.font.size = Pt(9); r.font.color.rgb = GRAY

spacer(doc, 4)

# Logo centered
add_image(doc, LOGO, 3.0)

spacer(doc, 2)
divider(doc)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("Security System Upgrade Proposal")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run("Mad Oak Bar & Yard  |  135 12th Street, Oakland, CA 94607")
r.font.size = Pt(12); r.font.color.rgb = TEAL

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_date.add_run("March 2026  |  Prepared by: Don Bucknor, COO")
r.font.size = Pt(10); r.font.color.rgb = GRAY; r.italic = True

divider(doc)
spacer(doc, 4)

# Description box
tbl = doc.add_table(rows=1, cols=1)
cell = tbl.rows[0].cells[0]
set_cell_shading(cell, "F2F7FB")
set_cell_border(cell, "003B71", "6")
p = cell.paragraphs[0]
r = p.add_run("Mad Oak Bar & Yard requires a complete security NVR upgrade to replace an end-of-life Korean-brand system with a modern 32-channel AI-capable platform. ASI 360 proposes the Luminys R54-32NA NVR with LumiSearch AI, reducing footage retrieval time from 4+ hours to under 2 minutes, while reusing all existing camera infrastructure and cabling.")
r.font.size = Pt(10)
p2 = cell.add_paragraph()
r2 = p2.add_run("Proposal Number: QUO202355  |  Valid through: April 5, 2026")
r2.font.size = Pt(9); r2.italic = True; r2.font.color.rgb = GRAY
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_header_bar(doc, "TABLE OF CONTENTS")
spacer(doc, 6)

toc_items = [
    ("Cover Letter", "1"),
    ("Title Page", "2"),
    ("Table of Contents", "3"),
    ("Needs Assessment", "4"),
    ("Scope of Work", "5"),
    ("Project Schedule & Milestones", "9"),
    ("Investment Summary", "11"),
    ("Risk Analysis", "13"),
    ("Project Deliverables", "14"),
    ("About ASI 360", "15"),
    ("Terms & Conditions", "16"),
    ("Authorization", "17"),
]

tbl_toc = doc.add_table(rows=len(toc_items), cols=2)
tbl_toc.style = 'Table Grid'

for ri, (name, pg) in enumerate(toc_items):
    row = tbl_toc.rows[ri]
    lc = row.cells[0]
    rc = row.cells[1]
    bg = "F2F7FB" if ri % 2 == 0 else "FFFFFF"
    set_cell_shading(lc, bg)
    set_cell_shading(rc, bg)
    set_cell_border(lc, "CCCCCC", "2")
    set_cell_border(rc, "CCCCCC", "2")

    lr = lc.paragraphs[0].add_run(name)
    lr.font.size = Pt(10)
    rr = rc.paragraphs[0].add_run(pg)
    rr.font.size = Pt(10)
    rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Quick Facts callout below TOC
spacer(doc, 10)
tbl_qf = doc.add_table(rows=1, cols=1)
cell_qf = tbl_qf.rows[0].cells[0]
set_cell_shading(cell_qf, "EEF4FB")
set_cell_border(cell_qf, "003B71", "6")
qf_title = cell_qf.paragraphs[0]
r_qft = qf_title.add_run("  PROPOSAL QUICK FACTS")
r_qft.bold = True; r_qft.font.size = Pt(10); r_qft.font.color.rgb = NAVY

facts = [
    ("Client", "Daniel Cukierman — Mad Oak Bar & Yard, Oakland, CA"),
    ("Proposal #", "QUO202355  |  Valid through April 5, 2026"),
    ("Option A", "$3,500 net  \u2022  Deposit: $2,913  \u2022  Balance: $587  \u2022  8TB / ~14-day retention"),
    ("Option B", "$3,750 net  \u2022  Deposit: $3,126  \u2022  Balance: $624  \u2022  16TB / ~30-day retention  \u2022  Recommended"),
    ("Hardware", "100% due with deposit  \u2022  Ships from TX, 5-day lead time"),
    ("Install Date", "Monday, March 23, 2026  (backup: March 25)"),
    ("Warranty", "90-day labor  \u2022  3-year NVR hardware  \u2022  1-year camera hardware"),
]
for label, val in facts:
    fp = cell_qf.add_paragraph()
    fr1 = fp.add_run(f"  {label}:  ")
    fr1.bold = True; fr1.font.size = Pt(9); fr1.font.color.rgb = TEAL
    fr2 = fp.add_run(val)
    fr2.font.size = Pt(9)
cell_qf.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — NEEDS ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_header_bar(doc, "NEEDS ASSESSMENT")
spacer(doc, 4)

body(doc, "ASI 360 has identified the following security infrastructure needs at Mad Oak Bar & Yard which are not currently being met by the existing Unix CCTV Magic View NVR system.")

spacer(doc, 4)
body(doc, "Needs:", bold=True)
needs = [
    ("Rapid Footage Retrieval", "Current system requires 4+ hours to locate and export footage after incidents. This creates critical liability exposure for police requests, insurance claims, and internal investigations."),
    ("Unified Camera Management", "22–24 cameras across 3 hardware generations (legacy 4K, mid-gen 2K, Dollartec IMAX series) are managed with no consistent IP scheme. Camera loss or misidentification is a constant operational risk."),
    ("Remote Access", "Owner and management cannot view live feeds or export footage remotely via phone or web browser. All review requires on-site presence."),
    ("AI-Assisted Search", "All footage review is entirely manual. No ability to search by person description, clothing color, or time-based behavior. Every incident review takes hours."),
    ("System Modernization", "The Unix CCTV Magic View NVR is discontinued, runs outdated firmware, and is incompatible with modern camera protocols required for future upgrades."),
]
for i, (title, desc) in enumerate(needs, 1):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '360'); ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '40')
    pPr.append(sp)
    r1 = p.add_run(f"{i}.  {title} \u2014 "); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(desc); r2.font.size = Pt(10)

spacer(doc, 4)
body(doc, "Market:", bold=True)
body(doc, "Mad Oak Bar & Yard \u2014 a licensed food and beverage venue in Oakland\u2019s Jack London Square district. The venue operates indoor bar, dining, and outdoor yard areas with 22\u201324 cameras covering all entry points, bar service areas, dining, and the outdoor yard. The venue is open to the public and requires continuous, reliable surveillance coverage for guest safety, staff accountability, and regulatory compliance.")

spacer(doc, 4)
body(doc, "Solution:", bold=True)
body(doc, "Replace the legacy NVR with the Luminys R54-32NA 32-channel AI platform. Install one Luminys N3T-8LA2 8MP AI turret camera at the dining area entry. Configure all 22\u201324 existing cameras under a unified static IP scheme using ONVIF/RTSP bridges for legacy camera compatibility. Enable LumiSearch AI on the new Luminys camera as Phase 1 of a fleet-wide AI transition over 18\u201324 months.")

spacer(doc, 4)
body(doc, "Sources:", bold=True)
body(doc, "\u2022 Site conversation and camera inventory review \u2014 Don Bucknor, ASI 360 COO, March 2026")
body(doc, "\u2022 ADI Distributor Quote #14395033 \u2014 Luminys R54-32NA, N3T-8LA2, N3T-4LA2 pricing confirmed")
body(doc, "\u2022 Luminys Product Catalog Volume 2, 2025 (EN-US) \u2014 technical specifications")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5+ — SCOPE OF WORK
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_header_bar(doc, "SCOPE OF WORK")
spacer(doc, 4)

body(doc, "Security is always of primary importance. Listed below are the specific measures ASI 360 will execute at Mad Oak Bar & Yard to upgrade the security video infrastructure.", size=10)
spacer(doc, 2)

# ── Section 1: NVR Hardware Swap ──
heading(doc, "1. NVR Hardware Swap", 2)
for b in [
    "Remove existing Unix CCTV Magic View NVR from rack",
    "Install new Luminys R54-32NA 32-channel AI NVR (H.265+, LumiSearch-ready)",
    "Retain existing PoE shelf, rack enclosure, and all cable runs \u2014 zero new wall work",
    "Configure NVR for remote access via web browser + mobile app (iOS/Android)",
    "Set up admin, manager, and view-only user accounts with secure credentials",
]:
    bullet(doc, b)

spacer(doc, 4)
product_card(doc,
    f"{IMG_DIR}/nvr_R5432NA_front.jpg", 2.0,
    "Luminys R54-32NA \u2014 32-Channel 8K NVR",
    "Your new system\u2019s brain \u2014 replaces the Unix CCTV Magic View",
    [
        "32 channels \u2014 supports all 22\u201324 cameras with room to grow",
        "8K / 4K / 1080p multi-format recording \u2014 H.265+ compression",
        "LumiSearch AI built in \u2014 search footage by person, clothing, object",
        "Dual Gigabit NIC \u2014 stable network throughput for all 32 channels",
        "3-year manufacturer hardware warranty",
    ],
    "Drops directly into your existing rack. Zero new cable runs \u2014 reuses all current infrastructure."
)
spacer(doc, 4)

# ── Section 2: Storage ──
heading(doc, "2. Hard Drive \u2014 Storage Options (Choose One)", 2)
body(doc, "Storage calculated for 22 cameras (14 \u00d7 4K @ avg 2.5 Mbps + 8 \u00d7 2K @ avg 1.2 Mbps, H.265+ compression, 24/7 recording):", bold=True)

tbl = doc.add_table(rows=4, cols=4)
tbl.style = 'Table Grid'
hdrs = ["Storage Option", "Capacity", "Estimated Retention", "Add-on Cost"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    set_cell_shading(cell, "003B71")
    r = cell.paragraphs[0].add_run(h)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE

rows_data = [
    ["Option A \u2014 2-Week Retention (Base)", "8 TB (1\u00d7 8TB HDD)", "~14\u201315 days @ 22 cams", "Included"],
    ["Option B \u2014 1-Month Retention (Recommended)", "16 TB (2\u00d7 8TB HDD)", "~30\u201332 days @ 22 cams", "+$185"],
    ["Option B \u2014 Alt (single drive)", "16 TB (1\u00d7 16TB HDD)", "~30\u201332 days @ 22 cams", "+$220"],
]
for ri, row_data in enumerate(rows_data):
    for ci, val in enumerate(row_data):
        cell = tbl.rows[ri+1].cells[ci]
        r = cell.paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        if ri % 2 == 0: set_cell_shading(cell, "F2F7FB")

note(doc, "Note: If camera count increases to 24, Option A retention drops to ~13 days. Option B still covers the full month. Hard drives are enterprise-grade surveillance rated (WD Purple / Seagate SkyHawk series).")

# ── Section 3: Static IP ──
spacer(doc, 4)
heading(doc, "3. Static IP Assignment & IP Table Routing", 2)
for b in [
    "Audit all existing cameras \u2014 document current IP, MAC address, and channel assignment",
    "Assign static IPs to all 22\u201324 cameras in sequential, logical scheme (e.g., 192.168.1.101\u2013.124)",
    "Update router/switch IP tables to lock MAC-to-IP bindings \u2014 IPs will never change on reboot",
    "Register each camera to its dedicated channel in correct sequential flow on NVR",
    "Document complete IP map (camera name \u2192 static IP \u2192 NVR channel \u2192 location label)",
    "Verify all cameras ping reliably and are visible from NVR and remote access",
]:
    bullet(doc, b)

# ── Section 4: Hybrid Camera Integration ──
spacer(doc, 4)
heading(doc, "4. Hybrid Camera Integration", 2)
for b in [
    "Map all existing cameras: legacy 4K (10-year-old), mid-gen 2K (5-year-old), Dollartec IMAX series, and any Luminys units",
    "Configure ONVIF/RTSP protocol bridges for non-native cameras on new NVR",
    "Set resolution, frame rate, and compression (H.265+ where supported, H.264 fallback for legacy)",
    "Verify timestamp sync via NTP on all cameras \u2014 correct timestamp displayed correctly on all channels",
    "Set up motion detection zones per camera appropriate to field of view",
]:
    bullet(doc, b)

# ── Section 5: New Camera ──
spacer(doc, 4)
heading(doc, "5. New Dining Area Camera (Luminys Brand)", 2)
for b in [
    "Supply and install 1\u00d7 Luminys 8MP AI turret camera at key entry point in dining area",
    "Mount, run cable to nearest PoE port, assign static IP, register to NVR channel",
    "Configure field of view, privacy masking if required, and AI detection zones",
    "This camera is LumiSearch AI native \u2014 will serve as the reference unit for future fleet transition",
]:
    bullet(doc, b)

spacer(doc, 4)
product_card(doc,
    f"{IMG_DIR}/cam_N3T8LA2_product.jpg", 2.0,
    "Luminys N3T-8LA2 \u2014 8MP AI Turret Camera",
    "Your first native LumiSearch AI camera \u2014 installed Day 1",
    [
        "8MP (4K) / True WDR \u2014 crisp image in mixed indoor/outdoor lighting",
        "LumiSearch AI native \u2014 AI footage search active from install day",
        "H.265+ compression \u2014 better quality at same storage cost",
        "IP67 weatherproof \u2014 indoor or outdoor installation",
        "100 ft (30.48m) IR night distance",
        "1-year manufacturer warranty",
    ],
    "Installed at the dining area entry. Each future Luminys camera added expands AI search across more of your fleet."
)
spacer(doc, 4)

# ── Section 6: LumiSearch ──
heading(doc, "6. LumiSearch AI \u2014 Phased Rollout Strategy", 2)
body(doc, "LumiSearch performs best as a unified fleet (all Luminys cameras), but partial benefit is available immediately on the new dining area camera and any future Luminys additions.", italic=True)
spacer(doc, 2)
for b in [
    "Phase 1 (This Install): LumiSearch enabled on 1\u00d7 new Luminys dining area camera \u2014 AI search active on Day 1",
    "Phase 2 (Ongoing): As legacy cameras are replaced, each new Luminys unit added to LumiSearch pool",
    "Full AI Search (Future): Once 80%+ of cameras are Luminys native, cross-camera search becomes fully operational",
    "AI Search allows staff to search by: person description, clothing color, object, time of entry \u2014 cuts incident response from hours to minutes",
]:
    bullet(doc, b)

spacer(doc, 4)
add_image(doc, f"{IMG_DIR}/lumisearch_ui.jpg", 5.5,
    caption="LumiSearch AI \u2014 search footage instantly by person, clothing color, or behavior description")
spacer(doc, 2)
note(doc, "Note: Facial recognition module is NOT included in this proposal scope. LumiSearch behavioral/object search is the active AI feature for this engagement.")

# ── Section 7: Camera Refurbishment ──
spacer(doc, 4)
heading(doc, "7. Camera Refurbishment (Separate Line Item)", 2)
for b in [
    "Inspect broken/offline camera \u2014 diagnose: physical damage, cable failure, or board failure",
    "Provide repair estimate or replacement recommendation after on-site diagnosis",
    "This is scoped and priced separately pending diagnosis \u2014 not included in base total",
]:
    bullet(doc, b)

# ── Section 8: Remote Access + Training ──
spacer(doc, 4)
heading(doc, "8. Remote Access, Notifications & Training", 2)
for b in [
    "Configure remote access: web browser (LAN/WAN) + mobile app (iOS/Android) for Daniel and designated staff",
    "Set up motion alert notifications pushed to phone (email + push notification)",
    "Configure alert zones: high-priority areas trigger immediate notification, perimeter cameras on schedule",
    "30-minute staff training session: how to review footage, export clips, use mobile app, manage alerts",
    "Provide written IP map reference sheet and login credentials document",
]:
    bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
# PROJECT SCHEDULE
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_header_bar(doc, "PROJECT SCHEDULE & MILESTONES")
spacer(doc, 4)

body(doc, "The following 5-phase schedule maps the project management framework to specific deliverables and dates for the Mad Oak NVR upgrade. The ASI 360 Hub starts Phase 1 network discovery immediately \u2014 installation day is targeted for Monday, March 23, 2026 (backup: March 25).")
spacer(doc, 2)
note(doc, "\u26a0  NVR Lead Time: Luminys R54-32NA ships from distributor warehouse in Texas. Standard transit is 5 business days. Order placed March 9 \u2192 estimated arrival March 16. No rush shipping required for March 23 install date.")
spacer(doc, 4)

phases = [
    ("Phase 1: Conception & Initiation  |  March 9, 2026", [
        ("1.1  Project Charter", "Contract executed, deposit received, project officially launched. Scope confirmed with Daniel."),
        ("1.2  Plan Review", "On-site or remote review of existing camera inventory, current IP scheme, rack layout, and cable runs. Hub pre-configuration begins."),
        ("1.3  Initiation", "ASI 360 Hub deployed to Mad Oak network. Automated discovery scan identifies all 22\u201324 cameras, documents IPs, MAC addresses, and signal quality."),
    ]),
    ("Phase 2: Definition & Planning  |  March 9\u201314, 2026", [
        ("2.1  Scope & Goal Setting", "Finalize camera count, confirm storage option (A or B), design static IP scheme (192.168.x.101\u2013.124)."),
        ("2.2  Budget Approval", "Hardware deposit confirms NVR order. Luminys R54-32NA shipped from Texas \u2014 5 business day lead time, estimated arrival March 16."),
        ("2.3  Work Breakdown Schedule", "Installation day confirmed: Monday, March 23, 2026. Backup date: Wednesday, March 25. Day-of task list finalized."),
        ("2.4  Project Schedule / Gantt", "NVR order tracked in real time via ASI 360 Hub. Client notified upon receipt and staging completion."),
        ("2.5  Communication Plan", "Client contact protocol established. Daniel receives status updates via SMS/email at each phase gate."),
        ("2.6  Risk Management", "Hybrid camera ONVIF/RTSP compatibility pre-verified. Contingency: March 25 backup install date if delivery delayed."),
    ]),
    ("Phase 3: Launch & Execution  |  March 16\u201322, 2026", [
        ("3.1  Status & Tracking", "NVR received at ASI 360 staging. Firmware updated, base configuration applied, Luminys N3T-8LA2 dining camera tested."),
        ("3.2  KPI Monitoring", "All 22\u201324 camera IPs pre-mapped using Hub discovery data. Static IP table built and staged \u2014 ready to apply on install day."),
        ("3.3  Quality Assurance", "ONVIF/RTSP protocol compatibility verified for each camera model. Legacy 4K, IMAX series, and new Luminys unit all confirmed."),
        ("3.4  Progress Forecasts", "Final day-of schedule confirmed with Daniel. Access arrangements, parking, and key contacts confirmed for March 23."),
    ]),
    ("Phase 4: Performance & Control  |  March 23, 2026 \u2014 Installation Day", [
        ("4.1  Objective Execution", "Remove Unix CCTV Magic View NVR from rack. Install Luminys R54-32NA. Apply pre-staged IP table \u2014 all cameras registered to sequential channels."),
        ("4.2  Quality Deliverables", "All 22\u201324 cameras verified live on new NVR. Timestamps synced via NTP. Remote access (web + mobile) confirmed working from off-site."),
        ("4.3  Effort & Cost Tracking", "Luminys N3T-8LA2 dining area camera installed, cabled, and registered. LumiSearch AI Phase 1 activated on dining camera."),
        ("4.4  Performance Review", "Motion alert zones configured. 30-minute staff training with Daniel: footage review, clip export, mobile app, alert management."),
    ]),
    ("Phase 5: Project Close  |  March 23\u201325, 2026", [
        ("5.1  Postmortem", "Broken camera on-site diagnosis completed. Repair scope and cost documented. Separate repair quote issued if applicable."),
        ("5.2  Project Punchlist", "IP map reference sheet delivered. Credentials document (all logins, IPs, channel assignments) handed off to Daniel."),
        ("5.3  Final Report", "System sign-off executed. 90-day labor warranty activated. Balance invoice issued. Project closed in VTiger and Airtable."),
    ]),
]

for phase_title, tasks in phases:
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '20')
    pPr.append(sp)
    r = p.add_run(phase_title)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
    for task_name, task_desc in tasks:
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '360')
        pPr.append(ind)
        sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '30')
        pPr.append(sp)
        r1 = p.add_run(f"{task_name}  "); r1.bold = True; r1.font.size = Pt(10)
        r2 = p.add_run(task_desc); r2.font.size = Pt(10)

# Fleet Roadmap — placed here at end of Schedule section
spacer(doc, 8)
heading(doc, "Luminys Fleet Transition Roadmap", 2)
body(doc, "The goal is full Luminys unification for complete LumiSearch AI capability. This happens in natural replacement cycles \u2014 no forced early swaps.")
spacer(doc, 2)

tbl_rm = doc.add_table(rows=5, cols=4)
tbl_rm.style = 'Table Grid'
for ci, h in enumerate(["Phase", "Action", "LumiSearch Coverage", "Timeline"]):
    cell = tbl_rm.rows[0].cells[ci]
    set_cell_shading(cell, "003B71")
    r = cell.paragraphs[0].add_run(h)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE

roadmap_data = [
    ("1", "NVR swap + 1\u00d7 Luminys dining camera (this proposal)", "1 camera AI-native", "Now"),
    ("2", "Replace 2\u20134 highest-priority cameras with Luminys 4K units", "3\u20135 cameras AI-native", "3\u20136 months"),
    ("3", "Replace remaining legacy cameras as they age out", "10+ cameras AI-native", "6\u201318 months"),
    ("4", "Full Luminys fleet \u2014 LumiSearch at maximum effectiveness", "All cameras", "18\u201324 months"),
]
for ri, row_data in enumerate(roadmap_data):
    for ci, val in enumerate(row_data):
        cell = tbl_rm.rows[ri+1].cells[ci]
        r = cell.paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        if ri % 2 == 0: set_cell_shading(cell, "F2F7FB")

# ══════════════════════════════════════════════════════════════════════════════
# INVESTMENT SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_header_bar(doc, "INVESTMENT SUMMARY")
spacer(doc, 4)

# ── Option A ──────────────────────────────────────────────────────────────────
investment_table(doc,
    opt_label="Option A \u2014 2-Week Storage (8TB)",
    opt_color="003B71",
    hw_items=[
        ("Luminys R54-32NA NVR, 32-Channel 1.5U Dual NIC, 4SATA, LumiSearch+", "1", "$1,242.00", "$1,242.00"),
        ("8TB WD Purple / Seagate SkyHawk Surveillance HDD", "1", "$175.00",   "$175.00"),
        ("Luminys N3T-8LA2, 8MP Mini Turret, LumiLuxSmart, IR + White-light",  "1", "$324.00",   "$324.00"),
        ("Cables, Connectors, Mounting Hardware (misc)",                        "1", "$85.00",    "$85.00"),
    ],
    labor_items=[
        ("Labor: NVR Swap, Rack Install, Cable Management",          "1", "$500.00", "$500.00"),
        ("Labor: IP Table Routing, Static IPs (22\u201324 cameras)", "1", "$800.00", "$800.00"),
        ("Labor: Hybrid Camera Config, ONVIF/RTSP Setup",            "1", "$499.00", "$499.00"),
        ("Labor: Remote Access, Notifications, Training",            "1", "$300.00", "$300.00"),
        ("Camera Refurbishment Diagnosis (on-site eval)",            "1",  "$75.00",  "$75.00"),
    ],
    hw_sub="$1,826.00",
    labor_sub="$2,174.00",
    gross="$4,000.00",
    discount="\u2212$500.00",
    net="$3,500.00",
    deposit_hw="$1,826.00",
    deposit_labor="$1,087.00",
    deposit_total="$2,913.00",
    balance_labor="$1,087.00",
    balance_discount="\u2212$500.00",
    balance_total="$587.00",
)

spacer(doc, 8)

# ── Option B ──────────────────────────────────────────────────────────────────
investment_table(doc,
    opt_label="Option B \u2014 1-Month Storage (16TB)  \u2014  Recommended",
    opt_color="1A6E1A",
    hw_items=[
        ("Luminys R54-32NA NVR, 32-Channel 1.5U Dual NIC, 4SATA, LumiSearch+", "1", "$1,242.00", "$1,242.00"),
        ("8TB WD Purple / Seagate SkyHawk Surveillance HDD \u00d7 2 (dual-drive)", "2", "$175.00", "$350.00"),
        ("Luminys N3T-8LA2, 8MP Mini Turret, LumiLuxSmart, IR + White-light",  "1", "$324.00",   "$324.00"),
        ("Cables, Connectors, Mounting Hardware (misc)",                        "1", "$85.00",    "$85.00"),
    ],
    labor_items=[
        ("Labor: NVR Swap, Rack Install, Cable Management",          "1", "$500.00", "$500.00"),
        ("Labor: IP Table Routing, Static IPs (22\u201324 cameras)", "1", "$800.00", "$800.00"),
        ("Labor: Hybrid Camera Config, ONVIF/RTSP Setup",            "1", "$499.00", "$499.00"),
        ("Labor: Remote Access, Notifications, Training",            "1", "$300.00", "$300.00"),
        ("Labor: Dual-Drive RAID Configuration & Verification",      "1",  "$75.00",  "$75.00"),
        ("Camera Refurbishment Diagnosis (on-site eval)",            "1",  "$75.00",  "$75.00"),
    ],
    hw_sub="$2,001.00",
    labor_sub="$2,249.00",
    gross="$4,250.00",
    discount="\u2212$500.00",
    net="$3,750.00",
    deposit_hw="$2,001.00",
    deposit_labor="$1,125.00",
    deposit_total="$3,126.00",
    balance_labor="$1,124.00",
    balance_discount="\u2212$500.00",
    balance_total="$624.00",
)

spacer(doc, 4)
body(doc, "Hardware is invoiced 100% with the deposit order. Labor is split 50% with the deposit and 50% on project completion. The $500 ASI 360 client discount is applied against the balance due at completion.", size=9, italic=True)

# ══════════════════════════════════════════════════════════════════════════════
# RISK ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_header_bar(doc, "RISK ANALYSIS")
spacer(doc, 4)

body(doc, "Our analysis of the requirements for this project identifies the following potential risks. This list is not exhaustive, but covers all primary risks identified during scope review.")
spacer(doc, 4)

risks = [
    ("Equipment Delivery Delay",
     "Distributor ships R54-32NA from Texas \u2014 5 business days standard. Supply chain disruption could push delivery past March 16.",
     "Order confirmed same day contract signed. Carrier tracking shared with client upon shipment.",
     "Backup install date: Wednesday, March 25, 2026. No impact to project close date."),
    ("Legacy Camera ONVIF Incompatibility",
     "Cameras >10 years old may not support modern RTSP/ONVIF streams required by new NVR.",
     "Pre-verify protocol support for each camera model during Phase 1 Hub scan.",
     "Manual MJPEG stream fallback or camera replacement quoted separately at $324/unit."),
    ("Network Infrastructure Gaps",
     "Client switch/router may not support PoE+ for new 8MP camera or full 32-channel throughput.",
     "Audit switch specs during Phase 1 Hub discovery scan.",
     "Additional PoE+ switch quoted separately if required (~$150\u2013$300 depending on port count)."),
    ("Scope Expansion",
     "Client requests additional cameras or scope changes mid-project.",
     "Change order process documented in contract. All additions require written approval.",
     "Additional camera: $150/camera add-on labor rate. Hardware quoted at current ADI pricing."),
    ("Camera Refurbishment Unknown Cost",
     "Broken camera diagnosis may reveal damage requiring full replacement.",
     "Flat $75 diagnosis fee included in proposal. Repair quoted separately after assessment.",
     "Replacement camera quoted if repair not viable. Does not delay primary NVR install."),
]

tbl = doc.add_table(rows=len(risks)+1, cols=4)
tbl.style = 'Table Grid'
for ci, h in enumerate(["Risk", "Analysis", "Resolution", "Contingency"]):
    cell = tbl.rows[0].cells[ci]
    set_cell_shading(cell, "003B71")
    r = cell.paragraphs[0].add_run(h)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE

for ri, (risk, analysis, resolution, contingency) in enumerate(risks):
    row = tbl.rows[ri+1]
    for ci, val in enumerate([risk, analysis, resolution, contingency]):
        r = row.cells[ci].paragraphs[0].add_run(val)
        r.font.size = Pt(8)
        if ri % 2 == 0: set_cell_shading(row.cells[ci], "F2F7FB")

# ══════════════════════════════════════════════════════════════════════════════
# DELIVERABLES
# ══════════════════════════════════════════════════════════════════════════════
spacer(doc, 8)
section_header_bar(doc, "PROJECT DELIVERABLES")
spacer(doc, 4)

body(doc, "The following is a complete list of all project deliverables upon project completion and client sign-off.")
spacer(doc, 2)

deliverables = [
    ("New NVR Installed & Operational", "Luminys R54-32NA live in rack \u2014 all 22\u201324 cameras streaming on correct channels"),
    ("Static IP Documentation", "Complete IP map: camera name \u2192 static IP \u2192 NVR channel \u2192 location label"),
    ("Credentials Document", "All usernames, passwords, and access levels for NVR, web access, and mobile app"),
    ("Remote Access Configured", "Web browser (LAN/WAN) + mobile app (iOS/Android) verified working off-site"),
    ("Motion Alert Notifications", "Alert zones configured \u2014 notifications active to Daniel\u2019s phone and email"),
    ("New Dining Area Camera", "Luminys N3T-8LA2 installed, cabled, registered, LumiSearch AI activated"),
    ("30-Minute Staff Training", "Footage review, clip export, mobile app, alert management \u2014 with Daniel"),
    ("90-Day Labor Warranty", "Covers all configured settings, IP assignments, and remote access setup"),
    ("Manufacturer Hardware Warranty", "Luminys R54-32NA: 3-year warranty registered. N3T-8LA2: 1-year warranty registered"),
    ("Work Order Closed", "Signed work order closed in VTiger. Project archived in ASI 360 Hub."),
]

for i, (title, desc) in enumerate(deliverables, 1):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '360'); ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '30')
    pPr.append(sp)
    r1 = p.add_run(f"{i}.  {title} \u2014 "); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(desc); r2.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# COMPANY HISTORY
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_header_bar(doc, "ABOUT ASI 360 \u2014 ALLIED SYSTEMS INTEGRATIONS")
spacer(doc, 4)

add_image(doc, LOGO, 2.0, align=WD_ALIGN_PARAGRAPH.LEFT)
spacer(doc, 2)

body(doc, "About Us", bold=True)
body(doc, "ASI 360 (Allied Systems Integrations) is an Oakland-based security systems integrator specializing in physical security infrastructure for small businesses, bars, restaurants, retail, and commercial venues throughout the Bay Area. We design, install, and support complete security ecosystems \u2014 from single-camera upgrades to full multi-site deployments \u2014 with a focus on practical, cost-effective solutions built on enterprise-grade hardware.")
spacer(doc, 2)
body(doc, "Our approach is always infrastructure-first: we audit what you have, preserve what works, and replace only what\u2019s holding you back. Every project closes with a complete IP map, credential document, and staff training session \u2014 no black-box systems that only we can service.")
spacer(doc, 4)

body(doc, "Services", bold=True)
for svc in [
    "Security NVR/DVR upgrades and hybrid camera integration",
    "IP camera installation (indoor/outdoor, fixed/PTZ, AI-capable)",
    "Static IP scheme design and full network documentation",
    "Remote access configuration (web + mobile) and alert setup",
    "LumiSearch AI deployment and phased fleet transitions",
    "Access control systems (Keri Systems \u2014 door readers, credentials, audit logs)",
    "ASI 360 Hub network monitoring \u2014 camera health, uptime alerts, remote diagnostics",
    "Service agreements and 90-day labor warranty on all installs",
]:
    bullet(doc, svc)

spacer(doc, 4)
body(doc, "Technology Partners", bold=True)
body(doc, "ASI 360 is an authorized reseller and integration partner for Luminys Systems (AI cameras and NVRs) and Keri Systems (access control). All hardware is sourced through ADI Global Distribution \u2014 the industry\u2019s leading security distributor \u2014 ensuring authentic, warranted equipment at distributor pricing.")

spacer(doc, 4)
body(doc, "How to Contact ASI 360", bold=True)
body(doc, "ASI 360 | Allied Systems Integrations")
body(doc, "Don Bucknor, COO/Founder")
body(doc, f"{PHONE}  |  don@asi360.co  |  asi360.co")
body(doc, "Oakland, CA")

# ══════════════════════════════════════════════════════════════════════════════
# TERMS & CONDITIONS
# ══════════════════════════════════════════════════════════════════════════════
spacer(doc, 8)
section_header_bar(doc, "TERMS & CONDITIONS")
spacer(doc, 4)

terms = [
    "Hardware costs are due 100% with the deposit to place the equipment order. Labor is billed 50% with the deposit and 50% upon project completion and client sign-off.",
    "The $500 ASI 360 client discount is applied against the labor balance due at project completion.",
    "90-day labor warranty on all configured settings, IP assignments, and remote access setup",
    "Hardware manufacturer warranty applies to NVR and cameras (1\u20133 years depending on model)",
    "Camera refurbishment is scoped separately \u2014 repair cost quoted after diagnosis",
    "Change orders require written approval before execution",
    "Client responsible for providing site access and electrical power",
    "ASI 360 not responsible for pre-existing infrastructure failures discovered during install",
    "This proposal is valid for 30 days from date of issue",
    "All pricing subject to distributor availability at time of order",
]
for t in terms:
    bullet(doc, t)

spacer(doc, 6)
body(doc, "ASI 360 Hub Monitoring (Optional Post-Project Service)", bold=True)
body(doc, "The ASI 360 Hub is a lightweight network agent deployed during Phase 1 discovery. After project close, optional ongoing monitoring is available:", size=10)
for hub_term in [
    "Camera health monitoring \u2014 uptime alerts and offline camera detection pushed to your phone",
    "Remote diagnostics \u2014 ASI 360 technician access via secure VPN for troubleshooting without an on-site visit",
    "Power cycle management and Grafana uptime dashboard view",
    "First 6 months included at no charge. Ongoing: $25/month (month-to-month, cancel anytime)",
    "Hub monitoring is separate from the 90-day labor warranty and does not replace manufacturer warranties",
]:
    bullet(doc, hub_term, size=9)

# ══════════════════════════════════════════════════════════════════════════════
# AUTHORIZATION / SIGNATURE
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_header_bar(doc, "AUTHORIZATION")
spacer(doc, 4)

body(doc, "By signing below, Client authorizes ASI 360 / Allied Systems Integrations to proceed with the work described in this proposal under the terms and conditions stated herein.")
spacer(doc, 4)

# Option selection
p_opt = doc.add_paragraph()
r1 = p_opt.add_run("Selected Option:  ")
r1.bold = True; r1.font.size = Pt(10)
r2 = p_opt.add_run("\u2610 Option A \u2014 Net Total: $3,500  (Deposit: $2,913 / Balance: $587)")
r2.font.size = Pt(10)

p_opt2 = doc.add_paragraph()
p_opt2.add_run("                         ")
r3 = p_opt2.add_run("\u2610 Option B \u2014 Net Total: $3,750  (Deposit: $3,126 / Balance: $624)")
r3.font.size = Pt(10)

spacer(doc, 6)

# Signature table
sig_tbl = doc.add_table(rows=3, cols=2)
sig_data = [
    ("Client Signature", "ASI 360 / Don Bucknor, COO"),
    ("Printed Name", "Don Bucknor"),
    ("Date", "Date"),
]
for ri, (left_label, right_label) in enumerate(sig_data):
    lc = sig_tbl.rows[ri].cells[0]
    rc = sig_tbl.rows[ri].cells[1]
    set_cell_border(lc, "003B71"); set_cell_border(rc, "003B71")

    lr = lc.paragraphs[0].add_run(left_label)
    lr.font.size = Pt(9); lr.font.color.rgb = TEAL; lr.bold = True
    rr = rc.paragraphs[0].add_run(right_label)
    rr.font.size = Pt(9); rr.font.color.rgb = TEAL; rr.bold = True

    lc.add_paragraph()
    rc.add_paragraph()

spacer(doc, 6)

# White text BoldSign tags
tag_p = doc.add_paragraph()
tag_run = tag_p.add_run("{{sign|1|*|Client Signature|sig_client}} {{date|1|*|Date Signed|date_client}} {{text|1|*|Printed Name|name_client}}")
tag_run.font.color.rgb = WHITE; tag_run.font.size = Pt(8)

tag_p2 = doc.add_paragraph()
tag_run2 = tag_p2.add_run("{{sign|2|*|ASI 360 Authorized Signature|sig_asi360}} {{date|2|*|Date|date_asi360}}")
tag_run2.font.color.rgb = WHITE; tag_run2.font.size = Pt(8)

spacer(doc, 4)
body(doc, f"Questions? Contact Don Bucknor  |  {PHONE}  |  don@asi360.co  |  asi360.co", size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT_DOCX)
print(f"Saved: {OUT_DOCX}")
